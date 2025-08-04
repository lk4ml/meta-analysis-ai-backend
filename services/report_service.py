import pandas as pd
import os
from typing import List, Dict, Any
from datetime import datetime
import json
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors

class ReportService:
    def __init__(self):
        self.reports_dir = "reports"
        os.makedirs(self.reports_dir, exist_ok=True)
    
    def generate_csv_report(self, question_data: Dict[str, Any], papers_data: List[Dict[str, Any]]) -> str:
        """
        Generate CSV report from extracted data
        """
        try:
            # Prepare data for CSV
            report_data = []
            
            for paper in papers_data:
                extracted_data = paper.get('extracted_data') or {}
                screening_data = paper.get('screening_json') or {}
                
                row = {
                    'PMID': paper.get('pmid', ''),
                    'Title': paper.get('title', ''),
                    'Authors': paper.get('authors', ''),
                    'Publication_Date': paper.get('publication_date', ''),
                    'DOI': paper.get('doi', ''),
                    'Score': paper.get('score', 0.0),
                    
                    # Screening results
                    'Study_Design_Screen': screening_data.get('study_design', ''),
                    'Intervention_Screen': screening_data.get('intervention', ''),
                    'Population_Screen': screening_data.get('population', ''),
                    'Outcomes_Screen': screening_data.get('outcomes', ''),
                    'Treatment_Characteristics_Screen': screening_data.get('treatment_characteristics', ''),
                    
                    # Extracted data
                    'Study_Design': extracted_data.get('study_design', ''),
                    'Patient_Characteristics': extracted_data.get('patient_characteristics', ''),
                    'Treatment_Characteristics': extracted_data.get('treatment_characteristics', ''),
                    'Intervention': extracted_data.get('intervention', ''),
                    'Comparison': extracted_data.get('comparison', ''),
                    'Outcomes': extracted_data.get('outcomes', ''),
                    
                    # Metadata
                    'Abstract': (paper.get('abstract', '') or '')[:500] + '...' if paper.get('abstract') else '',
                    'PDF_Link': paper.get('pdf_link', ''),
                    'MeSH_Terms': ', '.join(paper.get('mesh_terms', []) if isinstance(paper.get('mesh_terms'), list) else [])
                }
                report_data.append(row)
            
            # Create DataFrame
            df = pd.DataFrame(report_data)
            
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            question_id = question_data.get('id', 'unknown')
            filename = f"meta_analysis_report_{question_id}_{timestamp}.csv"
            filepath = os.path.join(self.reports_dir, filename)
            
            # Save CSV
            df.to_csv(filepath, index=False, encoding='utf-8')
            
            return filename
            
        except Exception as e:
            raise Exception(f"Error generating CSV report: {str(e)}")
    
    def generate_excel_report(self, question_data: Dict[str, Any], papers_data: List[Dict[str, Any]]) -> str:
        """
        Generate Excel report with multiple sheets
        """
        try:
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            question_id = question_data.get('id', 'unknown')
            filename = f"meta_analysis_report_{question_id}_{timestamp}.xlsx"
            filepath = os.path.join(self.reports_dir, filename)
            
            with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
                # Sheet 1: Summary
                summary_data = {
                    'Research Question': [question_data.get('original_text', '')],
                    'Rephrased Question': [question_data.get('rephrased_text', '')],
                    'Total Papers Found': [len(papers_data)],
                    'Average Score': [sum(p.get('score', 0) for p in papers_data) / len(papers_data) if papers_data else 0],
                    'Report Generated': [datetime.now().strftime("%Y-%m-%d %H:%M:%S")],
                    'PICO Criteria': [json.dumps(question_data.get('pico_json', {}), indent=2)]
                }
                summary_df = pd.DataFrame(summary_data)
                summary_df.to_excel(writer, sheet_name='Summary', index=False)
                
                # Sheet 2: Screening Results
                screening_data = []
                for paper in papers_data:
                    screening_json = paper.get('screening_json', {})
                    row = {
                        'PMID': paper.get('pmid', ''),
                        'Title': paper.get('title', '')[:100] + '...' if paper.get('title') else '',
                        'Score': paper.get('score', 0.0),
                        'Study_Design': screening_json.get('study_design', ''),
                        'Intervention': screening_json.get('intervention', ''),
                        'Population': screening_json.get('population', ''),
                        'Outcomes': screening_json.get('outcomes', ''),
                        'Treatment_Characteristics': screening_json.get('treatment_characteristics', '')
                    }
                    screening_data.append(row)
                
                screening_df = pd.DataFrame(screening_data)
                screening_df.to_excel(writer, sheet_name='Screening Results', index=False)
                
                # Sheet 3: Extracted Data
                extraction_data = []
                for paper in papers_data:
                    extracted_data = paper.get('extracted_data', {})
                    if extracted_data:  # Only include papers with extracted data
                        row = {
                            'PMID': paper.get('pmid', ''),
                            'Title': paper.get('title', '')[:100] + '...' if paper.get('title') else '',
                            'Authors': paper.get('authors', ''),
                            'Publication_Date': paper.get('publication_date', ''),
                            'Study_Design': extracted_data.get('study_design', ''),
                            'Patient_Characteristics': extracted_data.get('patient_characteristics', ''),
                            'Treatment_Characteristics': extracted_data.get('treatment_characteristics', ''),
                            'Intervention': extracted_data.get('intervention', ''),
                            'Comparison': extracted_data.get('comparison', ''),
                            'Outcomes': extracted_data.get('outcomes', '')
                        }
                        extraction_data.append(row)
                
                if extraction_data:
                    extraction_df = pd.DataFrame(extraction_data)
                    extraction_df.to_excel(writer, sheet_name='Extracted Data', index=False)
                
                # Sheet 4: Full Data
                full_data = []
                for paper in papers_data:
                    row = {
                        'PMID': paper.get('pmid', ''),
                        'Title': paper.get('title', ''),
                        'Authors': paper.get('authors', ''),
                        'Publication_Date': paper.get('publication_date', ''),
                        'DOI': paper.get('doi', ''),
                        'Score': paper.get('score', 0.0),
                        'Abstract': paper.get('abstract', ''),
                        'PDF_Link': paper.get('pdf_link', ''),
                        'MeSH_Terms': ', '.join(paper.get('mesh_terms', []) if isinstance(paper.get('mesh_terms'), list) else [])
                    }
                    full_data.append(row)
                
                full_df = pd.DataFrame(full_data)
                full_df.to_excel(writer, sheet_name='Full Paper Data', index=False)
            
            return filename
            
        except Exception as e:
            raise Exception(f"Error generating Excel report: {str(e)}")
    
    def get_report_path(self, filename: str) -> str:
        """Get full path to report file"""
        return os.path.join(self.reports_dir, filename)
    
    def list_reports(self) -> List[str]:
        """List all generated reports"""
        try:
            reports = []
            for filename in os.listdir(self.reports_dir):
                if filename.endswith(('.csv', '.xlsx', '.docx', '.pdf')):
                    filepath = os.path.join(self.reports_dir, filename)
                    stats = os.stat(filepath)
                    reports.append({
                        'filename': filename,
                        'size': stats.st_size,
                        'created': datetime.fromtimestamp(stats.st_ctime).strftime("%Y-%m-%d %H:%M:%S")
                    })
            return sorted(reports, key=lambda x: x['created'], reverse=True)
        except Exception as e:
            raise Exception(f"Error listing reports: {str(e)}")
    
    def delete_report(self, filename: str) -> bool:
        """Delete a report file"""
        try:
            filepath = os.path.join(self.reports_dir, filename)
            if os.path.exists(filepath):
                os.remove(filepath)
                return True
            return False
        except Exception as e:
            raise Exception(f"Error deleting report: {str(e)}")
    
    def generate_word_report(self, question_data: Dict[str, Any], papers_data: List[Dict[str, Any]]) -> str:
        """
        Generate Word (DOCX) report from extracted data
        """
        try:
            # Create a new document
            doc = Document()
            
            # Add title
            title = doc.add_heading('Meta-Analysis Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add metadata section
            doc.add_heading('Research Question', level=1)
            doc.add_paragraph(f"Original Question: {question_data.get('original_text', '')}")
            doc.add_paragraph(f"Rephrased Question: {question_data.get('rephrased_text', '')}")
            doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            
            # Add PICO criteria if available
            if question_data.get('pico_json'):
                doc.add_heading('PICO Criteria', level=2)
                pico = question_data.get('pico_json', {})
                doc.add_paragraph(f"• Population: {pico.get('population', 'Not specified')}")
                doc.add_paragraph(f"• Intervention: {pico.get('intervention', 'Not specified')}")
                doc.add_paragraph(f"• Comparison: {pico.get('comparison', 'Not specified')}")
                doc.add_paragraph(f"• Outcome: {pico.get('outcome', 'Not specified')}")
            
            # Add summary statistics
            doc.add_page_break()
            doc.add_heading('Summary Statistics', level=1)
            doc.add_paragraph(f"Total Papers Found: {len(papers_data)}")
            avg_score = sum(p.get('score', 0) for p in papers_data) / len(papers_data) if papers_data else 0
            doc.add_paragraph(f"Average Screening Score: {avg_score:.2f}")
            screened_count = len([p for p in papers_data if p.get('screening_json')])
            doc.add_paragraph(f"Papers Screened: {screened_count}")
            extracted_count = len([p for p in papers_data if p.get('extracted_data')])
            doc.add_paragraph(f"Papers with Extracted Data: {extracted_count}")
            
            # Add screening results section
            if any(p.get('screening_json') for p in papers_data):
                doc.add_page_break()
                doc.add_heading('Screening Results', level=1)
                
                # Create a table for screening results
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Table Grid'
                
                # Header row
                header_cells = table.rows[0].cells
                headers = ['PMID', 'Title', 'Score', 'Study Design', 'Intervention', 'Population', 'Outcomes']
                for i, header in enumerate(headers):
                    header_cells[i].text = header
                    # Bold the header
                    for paragraph in header_cells[i].paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                
                # Add data rows
                for paper in papers_data[:20]:  # Limit to first 20 papers
                    if paper.get('screening_json'):
                        row_cells = table.add_row().cells
                        screening = paper.get('screening_json', {})
                        
                        row_cells[0].text = paper.get('pmid', '')
                        row_cells[1].text = (paper.get('title', '')[:50] + '...') if paper.get('title') else ''
                        row_cells[2].text = f"{paper.get('score', 0):.1f}"
                        row_cells[3].text = screening.get('study_design', '')
                        row_cells[4].text = screening.get('intervention', '')
                        row_cells[5].text = screening.get('population', '')
                        row_cells[6].text = screening.get('outcomes', '')
            
            # Add extracted data section
            if any(p.get('extracted_data') for p in papers_data):
                doc.add_page_break()
                doc.add_heading('Extracted Data', level=1)
                
                for paper in papers_data:
                    if paper.get('extracted_data'):
                        extracted = paper.get('extracted_data', {})
                        
                        # Paper heading
                        doc.add_heading(f"PMID: {paper.get('pmid', '')}", level=2)
                        doc.add_paragraph(f"Title: {paper.get('title', '')}")
                        doc.add_paragraph(f"Authors: {paper.get('authors', '')}")
                        doc.add_paragraph(f"Publication Date: {paper.get('publication_date', '')}")
                        
                        # Extracted information
                        doc.add_heading('Study Details', level=3)
                        doc.add_paragraph(f"Study Design: {extracted.get('study_design', 'Not specified')}")
                        doc.add_paragraph(f"Patient Characteristics: {extracted.get('patient_characteristics', 'Not specified')}")
                        doc.add_paragraph(f"Treatment Characteristics: {extracted.get('treatment_characteristics', 'Not specified')}")
                        doc.add_paragraph(f"Intervention: {extracted.get('intervention', 'Not specified')}")
                        doc.add_paragraph(f"Comparison: {extracted.get('comparison', 'Not specified')}")
                        doc.add_paragraph(f"Outcomes: {extracted.get('outcomes', 'Not specified')}")
                        
                        doc.add_paragraph()  # Add spacing
            
            # Add references section
            doc.add_page_break()
            doc.add_heading('References', level=1)
            
            for i, paper in enumerate(papers_data, 1):
                if paper.get('title'):
                    ref_text = f"{i}. "
                    if paper.get('authors'):
                        ref_text += f"{paper['authors']}. "
                    ref_text += f"{paper['title']} "
                    if paper.get('publication_date'):
                        ref_text += f"({paper['publication_date']}). "
                    if paper.get('doi'):
                        ref_text += f"DOI: {paper['doi']}"
                    elif paper.get('pmid'):
                        ref_text += f"PMID: {paper['pmid']}"
                    
                    doc.add_paragraph(ref_text)
            
            # Generate filename and save
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            question_id = question_data.get('id', 'unknown')
            filename = f"meta_analysis_report_{question_id}_{timestamp}.docx"
            filepath = os.path.join(self.reports_dir, filename)
            
            doc.save(filepath)
            return filename
            
        except Exception as e:
            raise Exception(f"Error generating Word report: {str(e)}")
    
    def generate_pdf_report(self, question_data: Dict[str, Any], papers_data: List[Dict[str, Any]]) -> str:
        """
        Generate PDF report from extracted data
        """
        try:
            # Generate filename
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            question_id = question_data.get('id', 'unknown')
            filename = f"meta_analysis_report_{question_id}_{timestamp}.pdf"
            filepath = os.path.join(self.reports_dir, filename)
            
            # Create PDF document
            doc = SimpleDocTemplate(filepath, pagesize=letter,
                                  rightMargin=72, leftMargin=72,
                                  topMargin=72, bottomMargin=18)
            
            # Container for the 'Flowable' objects
            elements = []
            
            # Define styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=24,
                spaceAfter=30,
                alignment=1  # Center alignment
            )
            
            # Add title
            elements.append(Paragraph("Meta-Analysis Report", title_style))
            elements.append(Spacer(1, 0.2*inch))
            
            # Add metadata
            elements.append(Paragraph("Research Question", styles['Heading2']))
            elements.append(Paragraph(f"<b>Original Question:</b> {question_data.get('original_text', '')}", styles['Normal']))
            elements.append(Paragraph(f"<b>Rephrased Question:</b> {question_data.get('rephrased_text', '')}", styles['Normal']))
            elements.append(Paragraph(f"<b>Report Generated:</b> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
            elements.append(Spacer(1, 0.2*inch))
            
            # Add PICO criteria
            if question_data.get('pico_json'):
                elements.append(Paragraph("PICO Criteria", styles['Heading2']))
                pico = question_data.get('pico_json', {})
                pico_data = [
                    ['Population:', pico.get('population', 'Not specified')],
                    ['Intervention:', pico.get('intervention', 'Not specified')],
                    ['Comparison:', pico.get('comparison', 'Not specified')],
                    ['Outcome:', pico.get('outcome', 'Not specified')]
                ]
                pico_table = Table(pico_data, colWidths=[2*inch, 4*inch])
                pico_table.setStyle(TableStyle([
                    ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                    ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 10),
                    ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ]))
                elements.append(pico_table)
                elements.append(Spacer(1, 0.3*inch))
            
            # Add summary statistics
            elements.append(PageBreak())
            elements.append(Paragraph("Summary Statistics", styles['Heading2']))
            
            stats_data = [
                ['Total Papers Found:', str(len(papers_data))],
                ['Average Screening Score:', f"{sum(p.get('score', 0) for p in papers_data) / len(papers_data) if papers_data else 0:.2f}"],
                ['Papers Screened:', str(len([p for p in papers_data if p.get('screening_json')]))],
                ['Papers with Extracted Data:', str(len([p for p in papers_data if p.get('extracted_data')]))]
            ]
            
            stats_table = Table(stats_data, colWidths=[3*inch, 3*inch])
            stats_table.setStyle(TableStyle([
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
                ('GRID', (0, 0), (-1, -1), 1, colors.grey),
            ]))
            elements.append(stats_table)
            
            # Add screening results
            if any(p.get('screening_json') for p in papers_data):
                elements.append(PageBreak())
                elements.append(Paragraph("Screening Results", styles['Heading2']))
                elements.append(Spacer(1, 0.1*inch))
                
                # Create screening table
                screening_data = [['PMID', 'Score', 'Study Design', 'Intervention', 'Population']]
                
                for paper in papers_data[:15]:  # Limit to first 15 papers
                    if paper.get('screening_json'):
                        screening = paper.get('screening_json', {})
                        screening_data.append([
                            paper.get('pmid', '')[:10],
                            f"{paper.get('score', 0):.1f}",
                            screening.get('study_design', '')[:10],
                            screening.get('intervention', '')[:10],
                            screening.get('population', '')[:10]
                        ])
                
                screening_table = Table(screening_data, colWidths=[1.2*inch, 0.8*inch, 1.5*inch, 1.5*inch, 1.5*inch])
                screening_table.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                    ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                    ('GRID', (0, 0), (-1, -1), 1, colors.black)
                ]))
                elements.append(screening_table)
            
            # Add top extracted papers
            if any(p.get('extracted_data') for p in papers_data):
                elements.append(PageBreak())
                elements.append(Paragraph("Key Extracted Data", styles['Heading2']))
                
                count = 0
                for paper in papers_data:
                    if paper.get('extracted_data') and count < 5:  # Limit to top 5
                        count += 1
                        extracted = paper.get('extracted_data', {})
                        
                        elements.append(Spacer(1, 0.2*inch))
                        elements.append(Paragraph(f"<b>Paper {count} - PMID: {paper.get('pmid', '')}</b>", styles['Heading3']))
                        
                        # Create a condensed summary
                        summary_text = f"""
                        <b>Title:</b> {paper.get('title', '')[:100]}...<br/>
                        <b>Study Design:</b> {extracted.get('study_design', 'Not specified')[:100]}...<br/>
                        <b>Intervention:</b> {extracted.get('intervention', 'Not specified')[:100]}...<br/>
                        <b>Outcomes:</b> {extracted.get('outcomes', 'Not specified')[:100]}...
                        """
                        
                        elements.append(Paragraph(summary_text, styles['Normal']))
            
            # Build PDF
            doc.build(elements)
            return filename
            
        except Exception as e:
            raise Exception(f"Error generating PDF report: {str(e)}")